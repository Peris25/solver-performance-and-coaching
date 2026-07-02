"""Generate a coaching Word document for one solver.

This is a slimmed-down port of build_docs.py from the standalone skill.
It produces the same one-page report (with personal scorecard chart and
focused coaching) but lives in the webapp so reports can be downloaded
on demand from the dashboard.

The diagram generators are inlined here to avoid a dependency on the
external skill folder.
"""
from __future__ import annotations
import io
import re
from typing import Optional

import matplotlib
matplotlib.use("Agg")  # critical: no display backend in a server
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm


# Colors — aligned with the dashboard's red/white/black brand:
# - BLACK = "strong / on target" (was green in earlier versions)
# - RED = "needs work" (Solvit brand red, also signals attention)
# - WARN (muted red) = "close / over but not extreme"
NAVY = RGBColor(0x0E, 0x0E, 0x10)     # near-black for titles
GREEN = RGBColor(0x0E, 0x0E, 0x10)    # black for "good" — repurposed; name kept for compatibility
RED = RGBColor(0xD9, 0x25, 0x2A)      # Solvit brand red
AMBER = RGBColor(0xE9, 0x7A, 0x7C)    # muted red for "close"
GREY = RGBColor(0x5A, 0x5A, 0x60)

HEX_NAVY = "#0E0E10"
HEX_GREEN = "#0E0E10"   # black, see above
HEX_RED = "#D9252A"
HEX_AMBER = "#E97A7C"
HEX_GREY = "#5A5A60"


# --- formatters -----------------------------------------------------------

def fmt_hrs(h):
    if h is None:
        return "—"
    if h < 1:
        return f"{h*60:.0f} min"
    return f"{h:.1f} h"


def fmt_pct(p):
    return "—" if p is None else f"{p*100:.1f}%"


def fmt_rating(r):
    return "—" if r is None else f"{r:.2f} / 5"


def fmt_int(i):
    return "—" if i is None else f"{i:,}"


# --- diagram helpers ------------------------------------------------------

def _png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def personal_scorecard_png(stats: dict, team: dict, targets: dict) -> bytes:
    """Horizontal bars showing the solver's value against their target.

    Pure solver-vs-target. Team data is NOT shown here — the doc is about the
    solver's own performance only. (The webapp dashboard still shows team
    context, but in the doc the solver sees only their numbers vs targets.)
    """
    fig, ax = plt.subplots(figsize=(8, 3.6))

    def color_for(value, target, lower_is_better):
        if value is None: return HEX_GREY
        if lower_is_better:
            if value <= target: return HEX_GREEN          # black for on-target
            if value <= target * 1.5: return HEX_AMBER    # muted red for "close"
            return HEX_RED                                 # red for "needs work"
        else:
            if value >= target: return HEX_GREEN
            if value >= target * 0.75: return HEX_AMBER
            return HEX_RED

    total_tat = stats.get("avg_total_tat_hrs")
    rating = stats.get("avg_rating")
    sub_rate = stats.get("submission_rate")

    metrics = []
    if total_tat is not None:
        target = targets["total_tat_hours_max"]
        x_max = max(target * 2, total_tat * 1.2)
        metrics.append((
            "Average TAT (hrs)", total_tat, target, f"{total_tat:.1f} h", True,
            color_for(total_tat, target, True), x_max,
        ))
    if sub_rate is not None and stats.get("assigned_count", 0) >= 5:
        sv_pct = sub_rate * 100
        target_pct = targets["submission_rate_min"] * 100
        metrics.append((
            "Submission rate (%)", sv_pct, target_pct, f"{sv_pct:.0f}%", False,
            color_for(sub_rate, targets["submission_rate_min"], False), 100,
        ))
    if rating is not None and stats.get("n_ratings", 0) >= targets.get("min_ratings_for_judgement", 3):
        target = targets["rating_min"]
        metrics.append((
            "Client rating", rating, target, f"{rating:.2f}", False,
            color_for(rating, target, False), 5,
        ))

    if not metrics:
        ax.text(0.5, 0.5, "Not enough data to chart yet", ha="center", va="center",
                fontsize=11, color=HEX_GREY, style="italic", transform=ax.transAxes)
        ax.axis("off")
        return _png_bytes(fig)

    n = len(metrics)
    y_positions = list(range(n))[::-1]
    BAR_WIDTH_SCALE = 10  # everything plotted into a 0..10 x-range for consistency

    for y, (label, sv, tg, sv_fmt, lower_better, color, x_max) in zip(y_positions, metrics):
        scale = BAR_WIDTH_SCALE / x_max if x_max > 0 else 1

        # Background track (light grey) — shows the full scale
        ax.barh(y, BAR_WIDTH_SCALE, height=0.5, color="#F0F0F0", edgecolor="none", zorder=1)

        # Solver's bar (single colored bar — black if at target, red if not)
        ax.barh(y, sv * scale, height=0.5, color=color, edgecolor="none", zorder=2)

        # Solver's number label to the right of the bar
        ax.text(sv * scale + 0.15, y, sv_fmt, va="center", fontsize=10,
                fontweight="bold", color=color, zorder=4)

        # Target line — dashed vertical mark
        ax.plot([tg * scale, tg * scale], [y - 0.32, y + 0.32],
                color=HEX_NAVY, linewidth=2, linestyle="--", zorder=3)
        ax.text(tg * scale, y + 0.40, f"target {tg:g}", ha="center",
                fontsize=8, color=HEX_NAVY, style="italic", zorder=3)

        # Metric label on the left
        ax.text(-0.3, y, label, va="center", ha="right", fontsize=10,
                fontweight="bold", color=HEX_NAVY, zorder=3)

    ax.set_xlim(-3.5, BAR_WIDTH_SCALE + 2.5)
    ax.set_ylim(-0.6, n - 0.2)
    ax.set_yticks([])
    ax.set_xticks([])
    for spine in ("top", "right", "left", "bottom"):
        ax.spines[spine].set_visible(False)

    return _png_bytes(fig)


def time_two_clocks_png(stats: dict, targets: dict) -> bytes:
    """Plain-language visual of the solver's average TAT vs the 10h target.

    Big bar: total TAT (assignment → submission), the metric they're judged on.
    Small breakdown line below: response TAT and on-site TAT for context.
    """
    fig, ax = plt.subplots(figsize=(7.5, 2.8))

    total_tat = stats.get("avg_total_tat_hrs")
    rtat = stats.get("avg_response_tat_hrs")
    otat = stats.get("avg_onsite_tat_hrs")
    total_target = targets.get("total_tat_hours_max", 10.0)

    if total_tat is None:
        ax.text(0.5, 0.5, "Not enough completed jobs to chart yet",
                ha="center", va="center", fontsize=11, color=HEX_GREY,
                style="italic", transform=ax.transAxes)
        ax.axis("off")
        return _png_bytes(fig)

    # ---- Big bar: average total TAT ----
    scale_max = max(total_target * 3, total_tat * 1.2, total_target + 8)
    target_w = (total_target / scale_max) * 10
    y_main = 1.2

    # Target zone (black) and outside-target zone (faint red)
    ax.barh(y_main, target_w, height=0.6, left=0, color=HEX_GREEN,
            edgecolor="none", alpha=0.88)
    ax.barh(y_main, 10 - target_w, height=0.6, left=target_w, color=HEX_RED,
            edgecolor="none", alpha=0.28)

    # Inside-zone labels
    ax.text(target_w / 2, y_main, "✓ on target", ha="center", va="center",
            fontsize=10, color="white", fontweight="bold")
    ax.text(target_w + (10 - target_w) / 2, y_main, "over target",
            ha="center", va="center", fontsize=10, color=HEX_RED, fontweight="bold")

    # Target marker
    ax.plot([target_w, target_w], [y_main - 0.34, y_main + 0.34],
            color=HEX_NAVY, linewidth=2, linestyle="--")
    ax.text(target_w, y_main + 0.46, f"{int(total_target)}h target",
            ha="center", fontsize=9.5, color=HEX_NAVY, fontweight="bold")

    # Solver position
    value_x = min((total_tat / scale_max) * 10, 10)
    ax.scatter([value_x], [y_main - 0.46], marker="v", s=180,
               color=HEX_NAVY, zorder=4)
    ax.text(value_x, y_main - 0.74, f"you: {total_tat:.1f}h",
            ha="center", fontsize=10.5, color=HEX_NAVY, fontweight="bold")

    # Tick marks
    n_ticks = 5
    for i in range(n_ticks + 1):
        tick_x = (i / n_ticks) * 10
        tick_val = (i / n_ticks) * scale_max
        ax.plot([tick_x, tick_x], [y_main - 0.34, y_main - 0.40],
                color=HEX_GREY, linewidth=0.8)
        ax.text(tick_x, y_main - 1.00, f"{tick_val:.0f}h",
                ha="center", fontsize=8, color=HEX_GREY)

    # Left label
    ax.text(-0.3, y_main + 0.10, "Average time per job",
            ha="right", va="center", fontsize=11.5, fontweight="bold",
            color=HEX_NAVY)
    ax.text(-0.3, y_main - 0.18, "from assignment → submission",
            ha="right", va="center", fontsize=8.5,
            color=HEX_GREY, style="italic")

    # ---- Small breakdown line: response/on-site ----
    if rtat is not None and otat is not None:
        otat_min = otat * 60
        breakdown = (
            f"Of that {total_tat:.1f}h on average, "
            f"about {rtat:.1f}h is between scheduling and finishing, "
            f"and {otat_min:.0f} min is on-site with the vehicle."
        )
        ax.text(5, -0.5, breakdown, ha="center", va="center",
                fontsize=9, color=HEX_GREY, style="italic",
                transform=ax.transData)

    ax.set_xlim(-3.5, 12)
    ax.set_ylim(-1.5, 2.2)
    ax.axis("off")

    return _png_bytes(fig)


def submission_funnel_png(stats: dict, targets: dict) -> bytes:
    """Visual showing assigned -> valued with the pending leak.

    Two stacked horizontal bars:
      - Top: 'Assigned' (full grey bar = total jobs given)
      - Bottom: 'Submitted' (black portion = valued, red portion = pending leak)
    Plus the target line at 85%.
    """
    fig, ax = plt.subplots(figsize=(7.5, 2.6))

    assigned = stats.get("assigned_count", 0)
    valued = stats.get("valued_count", 0)
    pending = stats.get("pending_count", 0)
    rate = stats.get("submission_rate")
    target_rate = targets["submission_rate_min"]

    if assigned == 0:
        ax.text(0.5, 0.5, "No jobs assigned this period", ha="center", va="center",
                fontsize=11, color=HEX_GREY, style="italic", transform=ax.transAxes)
        ax.axis("off")
        return _png_bytes(fig)

    rate_color = HEX_GREEN if rate >= target_rate else (HEX_AMBER if rate >= 0.70 else HEX_RED)
    bar_full = 10  # plot into 0..10 x-range

    # Top bar — Assigned (full grey)
    ax.barh(1.3, bar_full, left=0, height=0.55, color="#E0E0E0",
            edgecolor="white", linewidth=2)
    ax.text(bar_full + 0.2, 1.3, f"{assigned}", va="center", ha="left",
            fontsize=11, fontweight="bold", color=HEX_NAVY)
    ax.text(-0.3, 1.45, "Assigned", ha="right", va="center",
            fontsize=11, fontweight="bold", color=HEX_NAVY)
    ax.text(-0.3, 1.10, "total jobs given\nto you", ha="right", va="center",
            fontsize=8.5, color=HEX_GREY, style="italic")

    # Bottom bar — Submitted (colored portion + pending leak)
    submitted_w = bar_full * (valued / assigned)
    pending_w = bar_full - submitted_w
    ax.barh(0.1, submitted_w, left=0, height=0.55, color=rate_color,
            edgecolor="white", linewidth=2)
    if pending > 0:
        ax.barh(0.1, pending_w, left=submitted_w, height=0.55, color="#F5C6C6",
                edgecolor="white", linewidth=2)
        # Pending count label
        if pending_w > 1.5:
            ax.text(submitted_w + pending_w / 2, 0.1, f"{pending} pending",
                    ha="center", va="center", fontsize=10, fontweight="bold",
                    color=HEX_RED)

    # Submitted count label inside the colored bar
    if submitted_w > 1.5:
        ax.text(submitted_w / 2, 0.1, f"{valued} valued",
                ha="center", va="center", fontsize=10, fontweight="bold", color="white")

    ax.text(-0.3, 0.25, "Submitted", ha="right", va="center",
            fontsize=11, fontweight="bold", color=HEX_NAVY)
    ax.text(-0.3, -0.10, "jobs you actually\nvalued", ha="right", va="center",
            fontsize=8.5, color=HEX_GREY, style="italic")

    # Target line at 85%
    target_x = bar_full * target_rate
    ax.plot([target_x, target_x], [-0.25, 0.45],
            color=HEX_NAVY, linewidth=2, linestyle="--")
    ax.text(target_x, 0.62, f"{int(target_rate*100)}% target",
            ha="center", fontsize=9, color=HEX_NAVY, fontweight="bold")

    # Rate summary at the bottom
    ax.text(bar_full / 2, -0.7, f"you: {valued}/{assigned} = {rate*100:.0f}%",
            ha="center", fontsize=12, fontweight="bold", color=rate_color)

    ax.set_xlim(-3.8, bar_full + 1.8)
    ax.set_ylim(-1.1, 2.0)
    ax.axis("off")

    return _png_bytes(fig)


# --- docx helpers ---------------------------------------------------------

def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _heading(doc, text, size=14, color=NAVY, space_before=12, space_after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def _para_with_inline(doc, text, size=11, space_after=6, italic=False, color=None):
    """Paragraph supporting **bold** and *italic* inline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*)")
    for part in pattern.split(text):
        if not part:
            continue
        run = p.add_run()
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run.text = part[1:-1]
            run.font.italic = True
        else:
            run.text = part
        run.font.size = Pt(size)
        run.font.italic = italic if not run.font.italic else True
        run.font.name = "Calibri"
        if color is not None:
            run.font.color.rgb = color
    return p


def _add_image(doc, png, width_in=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(io.BytesIO(png), width=Inches(width_in))


def _short_headline(stats, team, targets):
    vol = stats["volume"]
    rating = stats.get("avg_rating")
    n_ratings = stats.get("n_ratings", 0)
    rating_str = f"{rating:.2f}/5" if rating and n_ratings >= targets["min_ratings_for_judgement"] else None
    classifications = stats.get("classifications", {})
    needs = {k for k, v in classifications.items() if v == "needs_work"}

    time_flags = needs & {"response_tat", "onsite_tat"}
    submission_flags = needs & {"submission_rate"}
    rating_flag = "rating" in needs

    head = f"This month you completed **{vol}** valuations"
    if rating_str:
        head += f". Your clients rated you **{rating_str} out of 5**"
    head += "."

    if not needs:
        focus = (" You met every target this month. That's the result of steady habits, "
                 "not luck — keep doing what you're doing.")
    elif time_flags and submission_flags:
        focus = (" Two things to work on this month: **how long jobs take you** and "
                 "**finishing the jobs you accept**. Both are explained below in plain language.")
    elif time_flags and rating_flag:
        focus = (" Two things to look at this month: **how long jobs take you** "
                 "(explained below) and **how clients rated you**. Your team lead will "
                 "talk to you about the rating directly.")
    elif submission_flags and rating_flag:
        focus = (" Two things to look at this month: **finishing the jobs you accept** "
                 "(explained below) and **how clients rated you**. Your team lead will "
                 "talk to you about the rating directly.")
    elif time_flags:
        focus = (" The main thing to work on this month is **how long jobs take you** "
                 "from the moment you accept the job to when you submit it. "
                 "The steps below will help.")
    elif submission_flags:
        focus = (" The main thing to work on this month is **finishing the jobs you "
                 "accept**. The steps below will help.")
    elif rating_flag:
        focus = (" The one area to look at is **how clients rated you**. Your team "
                 "lead will talk to you directly — small changes usually move the rating up quickly.")
    else:
        focus = ""

    return head + focus


def _time_block(doc, stats, targets):
    _heading(doc, "Finishing jobs faster", size=13, space_before=10, space_after=4)

    total_target = targets.get("total_tat_hours_max", 10.0)
    avg_total = stats.get("avg_total_tat_hrs")
    avg_response = stats.get("avg_response_tat_hrs")
    avg_onsite = stats.get("avg_onsite_tat_hrs")

    if avg_total is not None:
        total_str = f"{avg_total:.1f} hours"
        gap_str = ""
        if avg_total > total_target:
            gap = avg_total - total_target
            gap_str = f" That's about **{gap:.1f} hours longer than the target**."
        _para_with_inline(doc,
            f"**The simple number:** From the moment a job is assigned to you, "
            f"to the moment you submit your valuation, you took on average "
            f"**{total_str}**. The target is **{int(total_target)} hours**.{gap_str}")

        # Breakdown of where the time went
        if avg_response is not None and avg_onsite is not None:
            avg_onsite_min = avg_onsite * 60
            _para_with_inline(doc,
                f"**Where that time goes:** about **{avg_response:.1f} hours** "
                f"between when you scheduled the job and when you submitted it, "
                f"of which **{avg_onsite_min:.0f} minutes** was actually on-site "
                f"with the vehicle. The rest is travel and admin time around the visit.")

    # Diagram: two clocks showing target zones and where the solver sits
    png = time_two_clocks_png(stats, targets)
    if png:
        _add_image(doc, png, width_in=5.8)

    _para_with_inline(doc, "**Three things that will close the gap:**")
    for i, line in enumerate([
        "**Accept or decline the job within 5 minutes** of getting the alert — don't let it sit.",
        "**Don't let jobs sit overnight.** If you can't get to it the same day, decline it so someone else can.",
        "**Same routine, every visit:** photos → exterior → interior → submit before leaving the site.",
    ], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f"{i}.  ")
        num.font.size = Pt(11)
        num.font.name = "Calibri"
        # Inline bold rendering for the rest
        pattern = re.compile(r"(\*\*[^*]+\*\*)")
        for part in pattern.split(line):
            if not part: continue
            r = p.add_run()
            if part.startswith("**") and part.endswith("**"):
                r.text = part[2:-2]
                r.font.bold = True
            else:
                r.text = part
            r.font.size = Pt(11)
            r.font.name = "Calibri"
    _para_with_inline(doc,
        "*One thing to try this week:* look at the slowest job you had last week. "
        "What held it up? That's usually the pattern worth fixing.",
        italic=False, color=GREY, size=11)


def _submission_block(doc, stats, targets):
    _heading(doc, "Finishing what you start", size=13, space_before=10, space_after=4)
    valued = stats.get("valued_count", 0)
    assigned = stats.get("assigned_count", 0)
    pending = stats.get("pending_count", 0)
    sub_pct_str = fmt_pct(stats.get("submission_rate"))

    _para_with_inline(doc,
        f"**The simple number:** This month you were given **{assigned} jobs**. "
        f"You completed **{valued}** of them ({sub_pct_str}). The other "
        f"**{pending} jobs** are still pending in your basket. The goal is to "
        f"close **at least 85%** of what you're given.")

    # Diagram: assigned -> submitted funnel with the pending gap
    png = submission_funnel_png(stats, targets)
    if png:
        _add_image(doc, png, width_in=5.8)

    pending_reasons = (stats.get("extra") or {}).get("top_pending_reasons") or []
    if pending_reasons:
        top_reason = pending_reasons[0]
        _para_with_inline(doc,
            f"Most of your pending jobs are sitting because of: "
            f"**{top_reason[0]}** ({top_reason[1]} jobs).")

    _para_with_inline(doc, "**Three things that will help you close more:**")
    for i, line in enumerate([
        "**Before driving to a client, call them first** to confirm they have the documents and the vehicle ready. This kills most 'No documents' pendings.",
        "If the client doesn't pick up the first time, **try again at a different time of day** before giving up — many people just miss the first call.",
        "When you mark a job as pending, **write the exact time you'll try again** (e.g. 'tomorrow 9 AM'). 'Later' or 'soon' never happens.",
    ], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f"{i}.  ")
        num.font.size = Pt(11)
        num.font.name = "Calibri"
        pattern = re.compile(r"(\*\*[^*]+\*\*)")
        for part in pattern.split(line):
            if not part: continue
            r = p.add_run()
            if part.startswith("**") and part.endswith("**"):
                r.text = part[2:-2]
                r.font.bold = True
            else:
                r.text = part
            r.font.size = Pt(11)
            r.font.name = "Calibri"
    _para_with_inline(doc,
        "*One thing to try this week:* for every job you mark pending, write down "
        "the next time you'll try. You'll close about half of them just from doing that.",
        color=GREY, size=11)


def _strong_block(doc):
    _heading(doc, "Thank you — keep going", size=13, space_before=10, space_after=4)
    _para_with_inline(doc,
        "You hit every target this month. That doesn't happen by accident — it's "
        "the result of habits you've built up, and it makes a real difference to "
        "the team and to our clients.")
    _para_with_inline(doc, "**Keep doing what works:**")
    _para_with_inline(doc,
        "The basics that got you here. Calling clients before traveling, taking "
        "your photos in the same order every time, presenting yourself well on "
        "site. Don't change a thing.")
    _para_with_inline(doc, "**Ready for more?** Here's where you can grow:")
    for line in [
        "**Take on harder jobs** — fleet vehicles, commercial trucks, post-accident assessments. The team needs people who can handle these.",
        "**Help a newer solver.** Twenty minutes on a call with someone struggling can save them weeks of figuring it out alone.",
        "**Tell us what you notice.** You see more cars than most people. If you spot a pattern with a channel partner, a vehicle type, or a location, flag it.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        pattern = re.compile(r"(\*\*[^*]+\*\*)")
        for part in pattern.split(line):
            if not part: continue
            r = p.add_run()
            if part.startswith("**") and part.endswith("**"):
                r.text = part[2:-2]
                r.font.bold = True
            else:
                r.text = part
            r.font.size = Pt(11)
            r.font.name = "Calibri"
    _para_with_inline(doc, "Thank you for the strong month.", color=GREY)


# --- main builder ---------------------------------------------------------

def _format_metric_delta(metric_key: str, m: dict) -> str:
    """Format a single metric's previous → current with direction arrow."""
    prev = m.get("previous")
    curr = m.get("current")
    direction = m.get("direction", "na")
    if prev is None or curr is None:
        return ""

    if metric_key == "total_tat":
        prev_s = f"{prev:.1f}h"; curr_s = f"{curr:.1f}h"
        label = "Average time per job"
    elif metric_key == "response_tat":
        prev_s = f"{prev:.1f}h"; curr_s = f"{curr:.1f}h"
        label = "Response time"
    elif metric_key == "onsite_tat":
        prev_s = f"{prev*60:.0f}min"; curr_s = f"{curr*60:.0f}min"
        label = "Time on site"
    elif metric_key == "submission_rate":
        prev_s = f"{prev*100:.0f}%"; curr_s = f"{curr*100:.0f}%"
        label = "Jobs finished"
    elif metric_key == "rating":
        prev_s = f"{prev:.2f}"; curr_s = f"{curr:.2f}"
        label = "Client rating"
    elif metric_key == "volume":
        prev_s = f"{int(prev)}"; curr_s = f"{int(curr)}"
        label = "Valuations done"
    elif metric_key == "assigned":
        prev_s = f"{int(prev)}"; curr_s = f"{int(curr)}"
        label = "Jobs assigned"
    else:
        prev_s = str(prev); curr_s = str(curr)
        label = metric_key

    arrow = "↑" if direction == "improved" else ("↓" if direction == "worsened" else "→")
    return f"**{label}**: {prev_s} {arrow} {curr_s}"


def _comparison_block(doc, comparison: dict, previous_period_label: str):
    """Render the vs-previous-period block.

    Shows the headline + 3-5 key metric deltas with arrows.
    """
    _heading(doc, f"How this compares to {previous_period_label}",
             size=13, space_before=10, space_after=4)

    headline = comparison.get("headline") or "broadly steady"
    # Capitalize only the first letter, preserve acronyms like "TAT"
    if headline:
        headline_display = headline[0].upper() + headline[1:]
    else:
        headline_display = "Broadly steady"
    _para_with_inline(doc, f"*{headline_display}.*", italic=False, color=GREY)

    # Show all non-NA metric deltas. total_tat first since that's the headline metric.
    metrics = comparison.get("metrics", {})
    interesting_keys = ["total_tat", "submission_rate", "volume", "assigned", "rating"]

    lines = []
    for key in interesting_keys:
        m = metrics.get(key, {})
        if m.get("direction") in ("improved", "worsened"):
            text = _format_metric_delta(key, m)
            if text:
                lines.append((text, m["direction"]))

    if not lines:
        _para_with_inline(doc, "All metrics broadly unchanged from the previous period.",
                          color=GREY, size=11)
        return

    # Limit to the top 4 most material changes
    for text, direction in lines[:4]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        pattern = re.compile(r"(\*\*[^*]+\*\*)")
        for part in pattern.split(text):
            if not part: continue
            r = p.add_run()
            if part.startswith("**") and part.endswith("**"):
                r.text = part[2:-2]
                r.font.bold = True
            else:
                r.text = part
            r.font.size = Pt(11)
            r.font.name = "Calibri"
            # Color the arrow's direction
            if direction == "improved":
                r.font.color.rgb = GREEN  # which is black in this brand
            elif direction == "worsened":
                r.font.color.rgb = RED


def build_doc(stats: dict, team: dict, targets: dict,
              period_label: str, personalised_intro: Optional[str] = None,
              comparison: Optional[dict] = None,
              previous_period_label: Optional[str] = None) -> bytes:
    """Render a one-page coaching docx and return its bytes.

    If `comparison` is provided (a dict from compare_snapshots), a
    "How this compares to <previous>" section is included.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    p = doc.add_paragraph()
    r = p.add_run("Solver Performance & Coaching")
    r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = NAVY; r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{stats['name']}  ·  {period_label}")
    r.font.size = Pt(12); r.font.color.rgb = GREY; r.font.name = "Calibri"

    # Talent-grid label (if available) — small italic sub-line under name
    tg = stats.get("talent_grid") or stats.get("extra", {}).get("talent_grid")
    if tg and tg.get("label"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"Talent grid: {tg['label']}")
        r.font.size = Pt(10); r.font.italic = True
        r.font.color.rgb = GREY; r.font.name = "Calibri"

    # Intro
    intro_text = personalised_intro or _short_headline(stats, team, targets)
    _para_with_inline(doc, intro_text, space_after=4)

    # Personal scorecard chart
    png = personal_scorecard_png(stats, team, targets)
    if png:
        _add_image(doc, png, width_in=6.0)

    # vs previous period — if comparison data is supplied
    if comparison and previous_period_label:
        _comparison_block(doc, comparison, previous_period_label)

    # Coaching blocks based on training_modules
    modules = stats.get("training_modules") or []
    classifications = stats.get("classifications", {})

    for mod in modules:
        if mod == "time":
            _time_block(doc, stats, targets)
        elif mod == "submission":
            _submission_block(doc, stats, targets)
        elif mod == "strong":
            _strong_block(doc)

    # If only rating is flagged (no coaching modules), explain that the team
    # lead will follow up rather than leaving the doc empty.
    if not modules and classifications.get("rating") == "needs_work":
        _heading(doc, "About your client rating", size=13, space_before=10, space_after=4)
        _para_with_inline(doc,
            f"Your clients rated you **{fmt_rating(stats['avg_rating'])} out of 5** this "
            f"month, from {stats['n_ratings']} ratings. The team is aiming for **4.5 or "
            "higher**. This isn't about your workflow — it's about how you come across to clients.")
        _para_with_inline(doc,
            "**Your team lead will talk to you about this directly.** It's usually small "
            "things — how you greet the client, how you explain what you're doing, "
            "how you close the visit — and they tend to move the rating up within a "
            "few weeks once you focus on them.")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
